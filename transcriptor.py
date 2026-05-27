# ============================================================================
# 1. INSTALACIÓN DE ENTORNO (Ejecutar primero en la misma celda)
# ============================================================================
# Se han fijado las versiones de pandas y numba para evitar conflictos de dependencia.
!pip install -q openai-whisper python-docx rapidfuzz
!pip install -q --upgrade "pandas==2.2.2" "numba<0.62.0"
!apt-get -qq update && apt-get -qq install -y ffmpeg > /dev/null 2>&1

# ============================================================================
# 2. IMPORTACIONES Y CONFIGURACIÓN
# ============================================================================
import os, re, sys, warnings
from datetime import datetime
# getpass ya no es necesario ya que Hugging Face y pyannote no se utilizan.
from google.colab import files
import torch
import whisper
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np

warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

USAR_DIARIZACION = False  # Cambiado a False según la solicitud del usuario (no usar Hugging Face/Pyannote)
NOMBRES_CANONICOS = ["Juan Pérez", "María García", "Pedro Rodríguez", "Ana López", "Carlos Díaz"]

# El bloque de Hugging Face ha sido eliminado ya que el usuario no desea usarlo.
# ============================================================================
# 3. FUNCIONES AUXILIARES
# ============================================================================
def corregir_nombres_fuzzy(texto, nombres_canonicos, umbral=85):
    """Reemplaza variaciones de nombres usando coincidencia difusa."""
    from rapidfuzz import process, fuzz
    # Dividimos preservando espacios y puntuación
    tokens = re.split(r'(\W+)', texto)
    for i, token in enumerate(tokens):
        if len(token) < 3 or not token[0].isupper():
            continue
        # Buscamos el nombre más parecido
        match, score, _ = process.extractOne(token, nombres_canonicos, scorer=fuzz.token_set_ratio)
        if score >= umbral:
            # Respetamos si el original iba en mayúsculas o título
            tokens[i] = match if token.isupper() else match.capitalize()
    return "".join(tokens)

def asignar_voces_a_parrafos(diarizacion, whisper_result):
    """Une diarización + transcripción palabra por palabra y agrupa por voz.
    Ahora gestiona la ausencia de diarización devolviendo una transcripción continua.
    """
    # Si no hay diarización o no hay 'words' en el resultado de whisper, devuelve el texto completo como un solo párrafo.
    # El `word_timestamps` en `transcribe` ahora es `False`, por lo que `whisper_result.get("words", [])` estará vacío.
    return [{"speaker": "TRANSCRIPCIÓN", "text": whisper_result["text"].strip(), "start": 0}]

def generar_word_formal(parrafos, nombre_audio, es_diarizacion=False):
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Inches(11), Inches(8.5)

    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Arial', Pt(11)

    titulo = doc.add_heading('ACTA DE SESIÓN LEGISLATIVA - DELTA AMACURO', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"Estado: Delta Amacuro\n")
    meta.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\n")
    meta.add_run(f"Audio: {nombre_audio}")

    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(6)
    sep.paragraph_format.space_after = Pt(12)
    sep.add_run("=" * 40)

    # Como `es_diarizacion` siempre será `False` aquí, siempre entraremos a este bloque.
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.5)
    if parrafos: # Asegurarse de que hay texto para añadir
        p.add_run(parrafos[0]["text"])

    doc.add_paragraph("\n")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, cargo in enumerate(["PRESIDENTE DEL CONSEJO LEGISLATIVO", "SECRETARIO DE CÁMARA"]):
        c = table.rows[0].cells[i].paragraphs[0]
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run("\n\n" + "_"*28 + f"\n{cargo}")

    nombre_word = nombre_audio.rsplit('.', 1)[0] + "_ACTA.docx"
    doc.save(nombre_word)
    return nombre_word

# ============================================================================
# 4. EJECUCIÓN PRINCIPAL
# ============================================================================
def transcriptor_legislativo_v2():
    print("📂 Sube el/los audio(s) de la sesión:")
    subidos = files.upload()
    if not subidos:
        print("⚠️ No se subieron archivos. Cancelando.")
        return

    if torch.cuda.is_available():
        mem = torch.cuda.get_device_properties(0).total_memory / 1e9
        modelo = "large-v3" if mem >= 14 else "medium"
        print(f"🤖 VRAM: {mem:.1f} GB → Modelo: {modelo.upper()}")
    else:
        modelo = "medium"
        print("⚠️ Sin GPU. Usando MEDIUM (más lento).")

    device = "cuda" if torch.cuda.is_available() else "cpu"
    whisper_model = whisper.load_model(modelo, device=device, download_root="/content/.whisper")

    # Diarización desactivada por el usuario, por lo que este bloque no es necesario.
    # if USAR_DIARIZACION:
    #     try:
    #         from pyannote.audio import Pipeline
    #         pipeline_diar = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1")
    #         print("✅ Pipeline de diarización cargado.")
    #     except Exception as e:
    #         print(f"⚠️ Falló pyannote: {e}. Continuando SIN diarización.")
    #         pipeline_diar = None

    for nombre_audio in subidos.keys():
        print(f"\n🎙️ Procesando: {nombre_audio}")
        torch.cuda.empty_cache()
        try:
            resultado = whisper_model.transcribe(
                nombre_audio,
                language="es",
                fp16=(device == "cuda"),
                beam_size=5,
                condition_on_previous_text=True,
                verbose=False,
                word_timestamps=False # Diarización y word_timestamps desactivados.
            )

            diarizacion = None # La diarización está permanentemente desactivada.
            # `asignar_voces_a_parrafos` ahora gestiona esto, devolviendo un único párrafo.
            parrafos = asignar_voces_a_parrafos(diarizacion, resultado)
            for p in parrafos:
                # corregir_nombres_fuzzy se aplica al texto completo ya que no hay diarización.
                p["text"] = corregir_nombres_fuzzy(p["text"], NOMBRES_CANONICOS)

            # `es_diarizacion` se pasa como `False` para asegurar el formato de transcripción continua.
            nombre_word = generar_word_formal(parrafos, nombre_audio, es_diarizacion=False)
            files.download(nombre_word)
            print(f"✅ Descargado: {nombre_word}")

        except Exception as e:
            print(f"❌ Error crítico en {nombre_audio}: {e}")
            import traceback; traceback.print_exc()

if __name__ == "__main__":
    transcriptor_legislativo_v2()
